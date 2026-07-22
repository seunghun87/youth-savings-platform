from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


SOURCE = Path(r"C:\side_pj\app_finacial\docs\reference\청년금융지원플랫폼 통합설계서 v1 1.docx")
TARGET = Path(r"C:\side_pj\app_finacial\docs\청년금융지원플랫폼 GitHub 통합설계서 v2.0.docx")


def replace_paragraph(paragraph, text):
    runs = paragraph.runs
    if runs:
        runs[0].text = text
        for run in runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def replace_cell(cell, text):
    lines = str(text).split("\n")
    paragraph = cell.paragraphs[0]
    runs = paragraph.runs
    if runs:
        runs[0].text = lines[0]
        for run in runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(lines[0])
    for line in lines[1:]:
        paragraph.add_run().add_break()
        paragraph.add_run(line)
    for extra in cell.paragraphs[1:]:
        replace_paragraph(extra, "")


def replace_table(table, rows):
    template = deepcopy(table.rows[1]._tr if len(table.rows) > 1 else table.rows[0]._tr)
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)
    for _ in rows[1:]:
        table._tbl.append(deepcopy(template))
    for row, values in zip(table.rows, rows):
        for cell, value in zip(row.cells, values):
            replace_cell(cell, value)


def main():
    doc = Document(SOURCE)
    p = doc.paragraphs
    replacements = {
        4: "Version 2.0  |  2026-07-17",
        5: "기준 산출물: GitHub main / af724b7 (React 18 · Vite · Capacitor · Express · Supabase)",
        11: "본 문서는 GitHub 저장소에 구현된 청년 금융지원 통합 플랫폼의 화면 구성, 기능 요구사항, 데이터 모델 및 인터페이스 규격을 정의한다. React/Vite 프런트엔드, Capacitor Android 패키지, Express API와 Supabase 마이그레이션을 기준으로 작성하며 개발·테스트·운영 전환의 공통 참조 문서로 사용한다.",
        13: "포함: 정보구조, 화면 흐름, 화면별 상세 설계, 기능 명세, 실제 데이터 모델, REST API, 추천·납입·중도해지 규칙, 디자인 시스템",
        14: "제외: 회원 인증 상세, 운영자 CMS, 푸시 인프라, 외부 금융기관 신청 완료 콜백, 보안 감사 및 운영 정책",
        21: "청년 사용자가 저축 목표를 세우고 금융감독원 및 정책 저축상품을 탐색·추천받아 가입상품과 납입 이력을 한 흐름에서 관리하는 Android 우선 모바일 서비스. 웹에서는 최대 430px 모바일 프레임으로 동일 경험을 제공한다.",
        25: "통합: 금융감독원 적금 데이터와 수동 등록 정책상품을 동일 상품 모델로 조회",
        26: "개인화: 월 납입액·기간·나이·연소득을 기준으로 자격을 필터링하고 추천 점수 산출",
        27: "실행 관리: 관심상품, 가입상품, 납입 이력, 목표 달성률, 만기·중도해지 상태를 영속 저장",
        28: "정합성: 정액·자유·증액 적립식의 납입 규칙과 월 한도를 API에서 검증",
        33: "기본 앱은 5탭(홈·찾기·플랜·혜택·MY) 구조다. 온보딩 완료 여부는 clientId별 Supabase 상태로 판단하며, 상품 상세·납입 기록·플랜 수정·해지 예상은 모달 또는 전용 뷰로 진입한다.",
        36: "[앱 실행]",
        37: "   └─ 사용자 상태 조회 GET /api/user-state/{clientId}",
        38: "        ├─ 미완료 → SCR-01 온보딩/목표설정 → 프로필·플랜 저장",
        39: "        └─ 완료   → SCR-02 홈",
        40: "                         │",
        41: "        ┌────────────────┼────────────────┬───────────────┐",
        42: "        ▼                ▼                ▼               ▼",
        43: "   SCR-03 찾기      SCR-05 플랜      SCR-06 혜택      SCR-07 MY",
        44: "        │                │                │               │",
        45: "   SCR-04 상품상세  납입기록/상품관리  혜택상세        프로필/설정",
        46: "        │                │                                │",
        47: "   관심상품·가입등록  목표수정·중도해지  ◄──────────────┘",
        48: "                         │",
        49: "                 상태 변경 후 전체 사용자 상태 재조회",
        52: "목적: 이름·나이·지역·연소득과 저축 목표·월 목표액·현재 보유액을 수집하여 사용자 상태를 최초 생성한다.",
        55: "목적: 기본 온보딩과 재설계에서 사용자 프로필 및 저축 플랜을 입력한다. 저장 성공 후 user_profile.onboarding_completed를 true로 갱신하고 통합 홈으로 이동한다.",
        59: "clientId는 현재 데모 디바이스 식별자를 사용하며 인증 계정으로의 전환이 필요하다.",
        60: "금액 입력은 원 단위 정수로 정규화하고 목표액·월 목표액은 0보다 커야 한다.",
        61: "저장 API 실패 시 화면 상태를 유지하고 오류 메시지를 제공한다.",
        62: "재진입 시 GET /api/user-state 응답을 initialState로 주입해 기존 값을 유지한다.",
        64: "목적: 목표 달성률, 현재 저축액, 최근 납입, 가입상품, 맞춤 추천과 정책 혜택 진입점을 한 화면에 요약한다.",
        66: "4.4 SCR-04 상품 찾기",
        67: "목적: 금융감독원·정책 저축상품을 검색하고 유형, 기간, 적립 방식 등으로 필터링하며 관심상품 저장과 상세 진입을 제공한다.",
        69: "4.5 SCR-05 청년 혜택",
        70: "목적: 일자리·주거·교육·금융복지문화 등 청년정책 콘텐츠를 탐색한다. 현재 혜택 데이터는 프런트 정적 데이터이며 운영 API 연동이 후속 과제다.",
        73: "목적: 금리, 기간, 월 납입 범위, 적립 방식과 추천 사유를 확인하고 관심 저장 또는 가입상품 등록으로 연결한다.",
        76: "목적: 프로필과 저장상품을 확인하고 사용자 설정 메뉴에 접근한다. 가입상품·납입·목표 관리는 플랜 탭에서 수행한다.",
        82: "FR-01~FR-21은 main 브랜치 구현 기준이다. 인증, 운영자 CMS, 혜택 API, 외부 신청 완료 연동은 11장의 후속 과제로 분리한다.",
        90: "6.4 DB 스키마 (구현 — Supabase PostgreSQL)",
        92: "user_input              추천 요청 원본(월 납입액, 기간, 나이, 연소득, 소득분위)",
        93: "savings_product         금융상품 마스터(금리 옵션, 연령·소득·기간·납입 한도, 적립 방식)",
        94: "recommendation          요청별 추천 결과 순위·예상금액·안내문",
        95: "user_profile            client_id별 이름·나이·지역·연소득·온보딩 완료 여부",
        96: "savings_plan            목표액·월 목표액·현재액",
        97: "savings_contribution    가입상품별 납입액·납입일; 변경 시 플랜 현재액 동기화",
        98: "saved_product           사용자 관심상품 연결(client_id, product_id)",
        99: "user_enrolled_product   가입상태, 시작·만기·종료일, 금리·약정액·기초잔액, 해지 정보",
        100: "RPC increment_plan_amount(client_id, amount)로 현재 저축액을 원자적으로 증감",
        101: "모든 사용자 상태 테이블은 client_id를 소유 키로 사용하며 service_role 백엔드를 통해 접근",
        102: "상품 테이블은 fin_co_no + fin_prdt_cd 자연 키로 금융감독원 동기화 중복을 방지",
        103: "RLS 활성화: anon 직접 접근은 정책 미정의로 거부하고 서버만 service_role로 접근",
        104: "금액은 원 단위 integer/bigint, 금리는 numeric, 상태·적립방식은 CHECK 제약으로 제한",
        106: "추천 결과는 상품에 정적으로 저장하지 않는다. 요청 조건으로 eligible 상품을 선필터한 뒤 금리·정책성·조건 적합도를 점수화하고 expected_amount와 checks를 응답 시점에 계산한다.",
        109: "months = ceil((target_amount - current_amount) / monthly_target)",
        110: "months = max(0, months), years = floor(months / 12), remainder = months % 12",
        111: "표시: 목표 달성률과 예상 완료 시점; 현재액이 목표액 이상이면 100%로 클램프",
        113: "current_amount = opening_balance 합계 + savings_contribution 합계",
        114: "progress_pct = min(100, current_amount / target_amount × 100)",
        115: "납입 생성·수정·삭제 시 increment_plan_amount RPC에 차액을 전달",
        117: "프런트는 사용자 상태 재조회 후 합산 값을 다시 표시한다. 서버가 가입상품 존재 및 적립방식별 허용액을 검증하므로 클라이언트 표시값만 신뢰하지 않는다.",
        119: "principal = monthly_amount × period_months",
        120: "monthly_rate = annual_rate / 100 / 12",
        121: "maturity = Σ(payment × (1 + monthly_rate × remaining_months))",
        123: "calculationService는 납입 회차별 단리 근사액을 합산하고 원 단위로 반올림한다. 실제 세후 수령액·우대금리 충족 여부는 금융기관 약관과 다를 수 있음을 notice로 고지한다.",
        126: "8. API 명세 (구현)",
        127: "현재 Express 라우트 기준이다. 성공 응답은 리소스 또는 배열을 직접 반환하며, 오류는 { error: message } 형태다. 인증은 아직 없고 clientId가 URL 경로의 소유 식별자로 사용된다.",
        131: "성공: 단건 객체, 배열 또는 204 No Content",
        132: "실패: { \"error\": \"사용자에게 표시 가능한 오류 메시지\" }",
        142: "모바일 뷰포트 최대 폭 430px, 데스크톱에서는 중앙 정렬된 모바일 프레임",
        143: "하단 5탭 고정 내비게이션, 모달은 배경 스크롤 잠금과 명확한 닫기 동작 제공",
        144: "카드 라운드 8~20px가 혼재하며 신규 공통 컴포넌트는 8px 이하로 정리 권장",
        145: "상태 변경 후 fetchUserSavingsState를 재호출해 서버 상태를 단일 진실 원천으로 유지",
    }
    for index, text in replacements.items():
        replace_paragraph(p[index], text)

    tables = doc.tables
    replace_table(tables[1], [
        ["버전", "일자", "작성자", "내용"],
        ["1.1", "2026-07-16", "PM", "초기 UI 프로토타입 역설계"],
        ["2.0", "2026-07-17", "Codex", "GitHub main af724b7 기준: 통합 저축 홈, Supabase 영속화, 추천·상품·납입·가입·중도해지 API와 최신 DB 반영"],
    ])
    replace_table(tables[3], [
        ["ID", "화면명", "라우트/뷰", "하단 탭", "진입 경로"],
        ["SCR-01", "온보딩", "onboarding", "숨김", "최초 실행"],
        ["SCR-02", "설계 입력", "PlanPrototype", "숨김", "온보딩/재설정"],
        ["SCR-03", "통합 홈", "home", "홈", "설계 완료 후 기본"],
        ["SCR-04", "상품 찾기", "find", "찾기", "하단 탭/홈 추천"],
        ["SCR-05", "청년 혜택", "benefits", "혜택", "하단 탭/홈"],
        ["SCR-06", "상품 상세", "Product detail", "유지", "상품 카드 선택"],
        ["SCR-07", "마이페이지", "my", "MY", "하단 탭"],
        ["SCR-08", "저축 플랜", "plan", "플랜", "하단 탭"],
        ["SCR-09", "납입/목표 수정", "Record/Goal modal", "해당 없음", "홈/플랜 CTA"],
        ["SCR-10", "가입/중도해지", "Enroll/Terminate modal", "해당 없음", "상품 상세/플랜"],
    ])
    replace_table(tables[11], [
        ["ID", "기능", "설명", "관련 화면", "P"],
        ["FR-01", "사용자 상태 초기화", "clientId 기준 프로필·플랜 기본 행 생성 및 조회", "SCR-01~07", "M"],
        ["FR-02", "온보딩 저장", "프로필과 목표 플랜 저장, 완료 상태 갱신", "SCR-01", "M"],
        ["FR-03", "목표 현황", "목표액·현재액·월 목표·달성률 표시", "SCR-02,05", "M"],
        ["FR-04", "상품 목록", "정책/시중 상품과 금리·기간·납입 조건 조회", "SCR-03", "M"],
        ["FR-05", "상품 동기화", "금융감독원 적금 API를 savings_product에 upsert", "운영 API", "S"],
        ["FR-06", "맞춤 추천", "자격 필터, 예상 만기액, 적합도 점수와 사유 반환", "SCR-02,03", "M"],
        ["FR-07", "관심상품", "상품 찜 추가·해제와 사용자 상태 반영", "SCR-03,04,07", "M"],
        ["FR-08", "가입상품 등록", "약정액·시작일·만기일·적립방식과 기초잔액 저장", "SCR-04,10", "M"],
        ["FR-09", "납입 생성", "가입상품에만 납입 허용, 플랜 현재액 원자적 증가", "SCR-08", "M"],
        ["FR-10", "납입 수정", "기존 금액과 차액만큼 플랜 현재액 조정", "SCR-08", "M"],
        ["FR-11", "납입 삭제", "기록 삭제 후 플랜 현재액 감소", "SCR-08", "M"],
        ["FR-12", "정액 적립 검증", "회차 납입액이 약정 monthly_amount와 동일해야 함", "SCR-08", "M"],
        ["FR-13", "증액 적립 검증", "납입 회차별 step 금액을 더한 예정액 검증", "SCR-08", "M"],
        ["FR-14", "자유 적립 검증", "동일 월 누적액이 상품 max_amount를 넘지 않도록 검증", "SCR-08", "M"],
        ["FR-15", "가입 상태 변경", "신청중·가입완료·만기완료·중도해지 상태 관리", "SCR-05,10", "M"],
        ["FR-16", "중도해지 예상", "납입 원금과 해지 사유·예상 지급액 확인 후 상태 저장", "SCR-10", "M"],
        ["FR-17", "프로필 수정", "이름·나이·지역·연소득 갱신", "SCR-07", "S"],
        ["FR-18", "정책 혜택 탐색", "혜택 카드와 상세 정보 탐색", "SCR-06", "S"],
        ["FR-19", "오프라인 폴백", "API 장애 시 추천/화면의 로컬 표시값으로 기능 유지", "SCR-02,03", "S"],
        ["FR-20", "입력 오류 처리", "서버 검증 메시지를 모달에 표시하고 입력값 유지", "SCR-08,10", "M"],
        ["FR-21", "Android 연결", "에뮬레이터에서 10.0.2.2:3000으로 API 호출", "전체", "M"],
    ])
    replace_table(tables[12], [
        ["필드", "타입", "설명", "기본값 / 제약"],
        ["client_id", "string", "현재 사용자/기기 식별자", "demo-device; 향후 auth uid"],
        ["name", "string", "사용자 이름", "빈 문자열 허용"],
        ["age", "int", "만 나이", "0 이상"],
        ["city", "string", "거주 지역", "시·도 문자열"],
        ["annual_income", "int", "연소득", "원 단위, 0 이상"],
        ["onboarding_completed", "boolean", "초기 설정 완료", "false"],
        ["target_amount", "bigint", "목표 저축액", "0 이상"],
        ["monthly_target", "bigint", "월 목표 납입액", "0 이상"],
        ["current_amount", "bigint", "현재 저축액", "RPC로 증감"],
    ])
    replace_table(tables[13], [
        ["필드", "타입", "설명"],
        ["id", "uuid", "상품 PK"],
        ["fin_co_no / fin_prdt_cd", "string?", "금융감독원 금융사·상품 자연 키"],
        ["name / bank", "string", "상품명과 금융기관"],
        ["base_rate / options", "numeric / jsonb", "기본금리와 기간별 금리 옵션"],
        ["product_type", "정책 | 시중", "상품 유형"],
        ["min_age / max_age", "int", "가입 연령 범위"],
        ["income_limit", "int?", "연소득 상한"],
        ["min_period / max_period", "int", "가입기간 범위(개월)"],
        ["monthly_limit", "int?", "월 납입 상한"],
        ["contribution_type", "fixed | flexible | step_up", "정액·자유·증액 적립식"],
        ["payment_frequency", "monthly | weekly | daily", "납입 주기"],
        ["min_monthly_amount", "int?", "최소 납입액"],
        ["installment_step_amount", "int?", "회차별 증액 금액"],
        ["source", "manual | finlife", "데이터 출처"],
    ])
    replace_table(tables[15], [
        ["화면", "필터", "정렬"],
        ["SCR-02 홈", "추천 결과 및 가입상품", "추천 rank, 최근 납입일"],
        ["SCR-03 찾기", "상품명·은행·유형·기간·적립방식", "기본금리/추천도"],
        ["SCR-05 플랜", "활성·종료 가입상품과 납입 이력", "가입일/납입일 DESC"],
        ["SCR-06 혜택", "청년정책 카테고리", "프런트 정의 순서"],
    ])
    replace_table(tables[16], [
        ["Method", "Endpoint", "설명", "주요 파라미터 / 응답"],
        ["GET", "/health", "서버 상태", "status, timestamp"],
        ["GET", "/api/products", "상품 목록", "SavingsProduct[]"],
        ["POST", "/api/products/sync", "금융감독원 적금 동기화", "upsert count"],
        ["POST", "/api/recommend", "추천 계산", "monthly_amount, period_months, age, personal_income"],
        ["GET", "/api/user-state/{id}", "통합 사용자 상태", "profile, plan, contributions, saved ids, enrolled products"],
        ["PUT", "/api/user-state/{id}/profile", "프로필 수정", "name, age, city, annual_income, onboarding_completed"],
        ["PUT", "/api/user-state/{id}/plan", "플랜 수정", "target_amount, monthly_target, current_amount"],
        ["POST", "/api/user-state/{id}/contributions", "납입 생성", "product_name, amount, contributed_at"],
        ["PUT", "/api/user-state/{id}/contributions/{contributionId}", "납입 수정", "product_name, amount, contributed_at"],
        ["DELETE", "/api/user-state/{id}/contributions/{contributionId}", "납입 삭제", "204 No Content"],
        ["POST", "/api/user-state/{id}/enrolled-products", "가입상품 등록", "상품·약정·기간·기초잔액"],
        ["PUT", "/api/user-state/{id}/enrolled-products/{productId}", "가입 상태/해지 갱신", "status, ended_at, termination_reason, payout"],
        ["PUT", "/api/user-state/{id}/saved/{productId}", "관심상품 토글", "saved boolean"],
    ])
    replace_table(tables[17], [
        ["HTTP", "코드", "상황"],
        ["200", "OK", "조회·수정 성공"],
        ["201", "CREATED", "납입 또는 가입상품 생성"],
        ["204", "NO_CONTENT", "납입 삭제"],
        ["400", "VALIDATION_ERROR", "금액·한도·적립 규칙·가입상품 조건 위반"],
        ["404", "NOT_FOUND", "라우트 또는 리소스 없음"],
        ["429", "RATE_LIMITED", "추천·동기화 요청 제한 초과"],
        ["500", "INTERNAL_ERROR", "Supabase/외부 API/서버 오류"],
    ])
    replace_table(tables[20], [
        ["컴포넌트", "역할"],
        ["Header", "화면 eyebrow·제목·알림 버튼"],
        ["BottomNav", "홈·찾기·플랜·혜택·MY 5탭 전환"],
        ["Product", "금리·기관·유형·찜 상태를 표시하는 상품 행"],
        ["PlanPage", "목표·가입상품·납입 현황과 편집 진입점"],
        ["TerminationPreview", "중도해지 원금·예상 지급액·사유 확인"],
        ["PlanPrototype", "온보딩 및 프로필·목표 저장"],
        ["SavingsPlanV2Prototype", "가입상품 등록·플랜 시나리오 실험 뷰"],
        ["Sonner Toast", "저장·실패·상태 변경 피드백"],
    ])
    replace_table(tables[21], [
        ["항목", "요구사항"],
        ["플랫폼", "React 18 + Vite 웹, Capacitor Android 패키징"],
        ["성능", "상품·상태 API는 모바일 네트워크에서 체감 지연을 줄이고 로딩 상태를 노출"],
        ["보안", "Supabase service_role 키는 백엔드 환경변수에만 보관; 프런트 노출 금지"],
        ["접근성", "아이콘 버튼 aria-label, 키보드 포커스, 충분한 대비와 44px 터치 영역"],
        ["정합성", "금액은 정수 검증, 납입 변경은 RPC 차액 반영, 상품 적립 규칙 서버 검증"],
        ["가용성", "추천 API 실패 시 로컬 폴백 또는 명시적 오류·재시도 상태 제공"],
        ["데이터", "금융감독원 동기화는 rate limit 적용, 자연 키 upsert, 출처 필드 유지"],
        ["개인정보", "clientId 임시 식별을 인증 uid로 교체하고 행 단위 소유권 정책 적용"],
    ])
    replace_table(tables[22], [
        ["구분", "항목", "비고"],
        ["인증", "회원가입·로그인·세션·계정복구", "clientId 데모 식별자 대체 필요"],
        ["보안", "사용자별 RLS 정책", "현재 service_role 서버 경유 구조에 인증 소유권 결합"],
        ["혜택", "청년정책 운영 API/CMS", "현재 프런트 정적 카드"],
        ["신청", "금융기관 딥링크와 신청 완료 콜백", "가입상품 등록은 사용자 입력 기반"],
        ["알림", "납입일·만기·마감 푸시", "FCM/APNs 및 사용자 동의 필요"],
        ["테스트", "API 통합·DB 마이그레이션·핵심 플로우 E2E", "현재 자동화 테스트 부족"],
        ["관측", "구조화 로그·오류 추적·헬스 대시보드", "운영 배포 전 필수"],
        ["재무", "세후 이자·우대금리·중도해지이율 정밀 계산", "현재 단리 근사/사용자 예상값"],
        ["UX", "로딩·빈 상태·오프라인·재시도 일관화", "화면별 처리 편차 정리"],
    ])

    xml_replacements = {
        "6.4 DB 스키마 (제안 — PostgreSQL)": "6.4 DB 스키마 (구현 — Supabase PostgreSQL)",
        "8. API 명세 (제안)": "8. API 명세 (구현)",
        "통합 설계서 v1.1": "GitHub 통합 설계서 v2.0",
        "4.4 SCR-04 저축 상품": "4.4 SCR-04 상품 찾기",
        "4.5 SCR-05 혜택 탐색": "4.5 SCR-05 청년 혜택",
    }
    parts = [doc.part, *(section.footer.part for section in doc.sections)]
    for part in parts:
        for node in part.element.iter(qn("w:t")):
            if node.text:
                for old, new in xml_replacements.items():
                    node.text = node.text.replace(old, new)

    TARGET.parent.mkdir(parents=True, exist_ok=True)
    doc.save(TARGET)


if __name__ == "__main__":
    main()
