from pathlib import Path
import json
import sys

from docx import Document


def main() -> None:
    source = Path(sys.argv[1])
    target = Path(sys.argv[2])
    doc = Document(source)
    data = {
        "paragraphs": [
            {"index": i, "style": p.style.name, "text": p.text}
            for i, p in enumerate(doc.paragraphs)
            if p.text.strip()
        ],
        "tables": [
            {
                "index": i,
                "rows": [[cell.text for cell in row.cells] for row in table.rows],
            }
            for i, table in enumerate(doc.tables)
        ],
    }
    target.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


if __name__ == "__main__":
    main()
