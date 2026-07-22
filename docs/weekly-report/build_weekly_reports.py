from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parent
TEMPLATE = ROOT / "weekly-report-template.docx"

WEEKS = [
    {
        "week": 1,
        "period": "2026. 6. 30. ~ 2026. 7. 5.",
        "date": "2026년  7월  5일",
        "goal": "청년 금융지원 통합 플랫폼의 핵심 문제를 정의하고, 금융상품 데이터 수집·추천을 위한 백엔드와 데이터베이스 기반을 구축한다.",
        "content": (
            "1. 저축 목표 설정-상품 탐색-맞춤 추천-납입 관리로 서비스 범위를 정의하였다.\n"
            "2. Express 서버, 환경변수, 오류 처리와 요청 제한 미들웨어를 구성하였다.\n"
            "3. Supabase 사용자 입력·금융상품·추천 결과 스키마와 마이그레이션을 작성하였다.\n"
            "4. 금융감독원 API의 적금 상품·기간별 금리 옵션 동기화를 구현하였다.\n"
            "5. 월 납입액·기간·나이·소득 기반 추천과 만기액 계산식을 구현·보정하였다."
        ),
        "method": "6/30 요구사항·화면 흐름 분석 2시간 / 7/1 DB·API 구조 구현 3시간 / 7/2 추천·계산 로직 보정 3시간 / 7/3 Supabase 스키마 점검 2시간 / 7/4 상품 동기화·예외 처리 확인 2시간 (총 12시간)",
        "result": "백엔드·DB·상품 동기화·추천 계산의 기반을 확보하였다. 금리 옵션 upsert와 추천 이력 보존 문제를 수정했으며 목표달성도는 약 90%이다.",
        "refs": "금융감독원 금융상품 한눈에 Open API 문서, Supabase JavaScript/PostgreSQL 문서, Express 공식 문서, 저장소 커밋 e62719a·6ad93e9·795eb75",
        "next": "프런트엔드 시안을 실제 API와 연결하고, 사용자 온보딩·상품 목록·추천 결과 화면을 구현한다. 배포 환경에서 Supabase 연결과 API 오류 상태도 확인한다.",
    },
    {
        "week": 2,
        "period": "2026. 7. 6. ~ 2026. 7. 12.",
        "date": "2026년  7월 12일",
        "goal": "React/Vite 기반 모바일 우선 화면을 구현하고 백엔드·Supabase와 연결하여 온보딩부터 상품 조회까지의 사용자 흐름을 완성한다.",
        "content": (
            "1. React 18·Vite에 모바일 디자인 시스템과 공통 UI를 구성하였다.\n"
            "2. 이름·지역·소득·자산·저축 목표를 받는 단계형 온보딩을 구현하였다.\n"
            "3. 상품 목록·검색·필터·상세·추천·홈·마이 화면을 제작하였다.\n"
            "4. 프런트 API 모듈로 Express 및 Supabase 실데이터를 연결하였다.\n"
            "5. Capacitor Android 패키징과 Render 배포·리전 설정을 정리하였다."
        ),
        "method": "7/6 배포 환경·리전 설정 2시간 / 7/7 UI 구조·디자인 토큰 2시간 / 7/8 온보딩 구현 3시간 / 7/9 API·Supabase 연결 및 Android 패키징 3시간 / 7/10 주요 화면 흐름 점검 2시간 (총 12시간)",
        "result": "온보딩-추천-상품 탐색이 실데이터로 연결되고 Android 패키징 구조를 확보하였다. 목표달성도는 약 90%이며 인증·운영 오류 처리는 후속 과제로 남겼다.",
        "refs": "React·Vite 공식 문서, Capacitor Android 문서, Supabase JavaScript 문서, Render Blueprint 문서, 저장소 커밋 5e1eccd·f1a6f56·fa56ef8",
        "next": "가입상품·납입 기록·저축 플랜을 하나의 사용자 상태로 통합하고, 정액·자유·증액 적립식 검증과 중도해지 반영까지 구현한다.",
    },
    {
        "week": 3,
        "period": "2026. 7. 13. ~ 2026. 7. 19.",
        "date": "2026년  7월 19일",
        "goal": "가입상품과 납입 이력을 반영하는 통합 저축 관리 프로토타입을 완성하고, 상품별 납입 규칙·목표 달성률·중도해지 계산의 정합성을 높인다.",
        "content": (
            "1. 홈·적금 찾기·내 플랜·혜택·MY의 5탭 모바일 UI를 통합하였다.\n"
            "2. 가입·관심상품, 납입 기록 CRUD, 플랜 수정 API와 화면을 연결하였다.\n"
            "3. 정액·자유·증액 적립식과 월/주 납입 규칙을 서버에서 검증하였다.\n"
            "4. 목표 달성률·상품별 기여도와 만기·중도해지 영향을 재계산하였다.\n"
            "5. 데이터 모델·REST API·계산식·화면 명세를 설계서로 정리하였다."
        ),
        "method": "7/13 사용자 상태·가입상품 모델 점검 2시간 / 7/14 납입 CRUD·검증 구현 3시간 / 7/15 플랜·중도해지 계산 개선 3시간 / 7/16 통합 UI·예외 상태 점검 2시간 / 7/17 설계서·빌드·Android 실행 확인 2시간 (총 12시간)",
        "result": "목표 설정부터 상품·납입·진행률·중도해지까지 이어지는 통합 프로토타입을 완성하였다. 서버 검증과 DB 마이그레이션으로 정합성을 높였고 목표달성도는 약 95%이다.",
        "refs": "프로젝트 통합설계서 v2.0, Supabase PostgreSQL/RPC 문서, Capacitor Android 문서, React 접근성 지침, 저장소 커밋 af724b7 및 migrations 004~010",
        "next": "실사용자 인증과 사용자별 RLS 정책을 적용하고, 핵심 API 통합 테스트·E2E 테스트·구조화 로그를 추가한다. 금융기관 신청 딥링크와 정책 혜택 운영 API 연동도 검토한다.",
    },
]


def set_run_font(run, size=8.5, bold=False):
    run.font.name = "함초롬돋움"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "함초롬돋움")
    run.font.size = Pt(size)
    run.bold = bold


def replace_cell(cell, text, size=8.5):
    p = cell.paragraphs[0]
    for r in list(p.runs):
        p._element.remove(r._element)
    for extra in list(cell.paragraphs[1:]):
        cell._tc.remove(extra._element)
    lines = text.split("\n")
    for i, line in enumerate(lines):
        if i:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run_font(run, size=size)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0


def replace_paragraph(paragraph, text, size=None):
    for r in list(paragraph.runs):
        paragraph._element.remove(r._element)
    run = paragraph.add_run(text)
    run.font.name = "함초롬돋움"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "함초롬돋움")
    if size:
        run.font.size = Pt(size)
    run.bold = True


def build(info):
    doc = Document(TEMPLATE)
    replace_paragraph(doc.paragraphs[2], f"2026-하계 세종창의학기제(집중이수제) 주간학습보고서 ({info['week']}주차)")
    replace_cell(doc.tables[0].cell(0, 1), "")
    t = doc.tables[1]
    replace_cell(t.cell(0, 1), "청년 금융지원 통합 플랫폼 앱 개발", 9)
    replace_cell(t.cell(1, 1), "")
    replace_cell(t.cell(1, 3), info["period"], 9)
    replace_cell(t.cell(2, 1), "")
    replace_cell(t.cell(2, 3), f"{info['week']}주차", 9)
    replace_cell(t.cell(2, 5), "12시간", 9)
    replace_cell(t.cell(3, 1), "")
    replace_cell(t.cell(3, 3), "자기주도창의전공 3", 9)
    replace_cell(t.cell(3, 5), "3학점", 9)
    replace_cell(t.cell(5, 1), info["goal"], 8.5)
    replace_cell(t.cell(6, 1), info["content"], 8.0)
    replace_cell(t.cell(7, 1), info["method"], 8.0)
    replace_cell(t.cell(8, 1), info["result"], 8.0)
    replace_cell(t.cell(9, 1), info["refs"], 7.8)
    replace_cell(t.cell(10, 1), info["next"], 8.0)
    replace_cell(t.cell(11, 3), info["date"], 9)
    out = ROOT / f"2026-하계_주간학습보고서_{info['week']}주차.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    for week in WEEKS:
        print(build(week))
